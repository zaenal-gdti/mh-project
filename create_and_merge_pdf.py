from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import pandas as pd
import numpy as np
import os
import shutil
from tqdm import tqdm
import glob
from zipfile import ZipFile 
from pypdf import PdfWriter
import os
from datetime import datetime
import shutil
import PyPDF2 as p2

def create_pdf_from_excel(df, dir_nm):
    if os.path.exists(dir_nm):
        shutil.rmtree(dir_nm)
    if not os.path.exists(dir_nm):
        os.makedirs(dir_nm)
    pd.options.mode.chained_assignment = None 

    for _, row in tqdm(df.iterrows(), total=df.shape[0]):
        data = row.copy()
        # Replace NaN with Empty String
        for i in data.keys():
            if str(data.get(i)) == 'nan':
                data[i] = ''
            else:
                data[i] = str(data[i])
        
        # Date Format
        if 'Tanggal' in data and data['Tanggal'] != '':
            data['Tanggal'] = pd.to_datetime(data['Tanggal']).strftime('%d-%m-%Y')
        
        # Open template
        with open('Template DASS-21 - AZ.docx', 'rb') as f:
            doc = Document(f)
        
        ## User Detail
        doc.paragraphs[1].text = doc.paragraphs[1].text.replace('<<Nama>>', data['Nama']).replace('<<Tanggal>>', data['Tanggal'])
        
        ## Score
        doc.tables[1].cell(1,1).text = doc.tables[1].cell(1,1).text.replace('<<SkorDepresi>>', data['SkorDepresi'])
        doc.tables[1].cell(1,2).text = doc.tables[1].cell(1,2).text.replace('<<InterpretasiDepresi>>', data['InterpretasiDepresi'])
        doc.tables[1].cell(2,1).text = doc.tables[1].cell(2,1).text.replace('<<SkorAnsietas>>', data['SkorAnsietas'])
        doc.tables[1].cell(2,2).text = doc.tables[1].cell(2,2).text.replace('<<InterpretasiAnsietas>>', data['Interpretasi Ansieta'])
        doc.tables[1].cell(3,1).text = doc.tables[1].cell(3,1).text.replace('<<SkorStres>>', data['SkorStres'])
        doc.tables[1].cell(3,2).text = doc.tables[1].cell(3,2).text.replace('<<InterpretasiStres>>', data['InterpretasiStres'])
        
        doc.tables[1].cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[1].cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[1].cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[1].cell(2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[1].cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[1].cell(3,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        nm = data['Nama']
        dob = data['DoB']
        tmp_docx = f'{dir_nm}/{nm}_{dob}_MHResult.docx'
        
        doc.save(tmp_docx)
        
        
        # Convert to PDF
        subprocess.run([
            'libreoffice', '--convert-to', 'pdf', tmp_docx, '--outdir', dir_nm
        ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        os.remove(tmp_docx)
        pdf_name = tmp_docx.replace('.docx', '.pdf')

def join_pdfs(mh_fls, mcu_fls):
    for i in tqdm(mh_fls):
        loc_files = [x for x in mcu_fls if os.path.basename(x).split('_')[:2] == os.path.basename(i).split('_')[:2]]
        try: 
            if loc_files:
                pdfs = [loc_files[0], i]
                merger = PdfWriter()

            for pdf in pdfs:
                merger.append(pdf)
            base_nm = os.path.basename(i)
            fold = '/'.join(i.split('/')[:-2]) + '/Result/' 
            if not os.path.exists(fold):
                os.makedirs(fold)
            res = fold +'tmp__'+base_nm
            if os.path.exists(res):
                os.remove(res)
            merger.write(res)
            merger.close()

            reader = p2.PdfReader(res)
            writer = p2.PdfWriter()



            passw = pd.to_datetime(base_nm.split('_')[1]).strftime('%d%m%Y') #.replace('.pdf', '')
            # new_fname = '_'.join(pdf_file.split('_')[:-1]) + '.pdf'
            # new_fname = out_path + os.path.basename(new_fname)
            for page in reader.pages:
                writer.add_page(page)

            # Set password (user password)
            writer.encrypt(user_password=passw)

            with open(res.replace('tmp__',''), "wb") as f:
                writer.write(f)   
            os.remove(res)

        except:
            if not loc_files:
                print("MCU file not found ==>,", base_nm)
            else:
                print("Error merging files for ==>", base_nm)
    